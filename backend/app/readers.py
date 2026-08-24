from pathlib import Path
from pypdf import PdfReader
from docx import Document

#File Reading Functions
def read_pdf(file_path:Path):
    pdf=PdfReader(file_path)
    text=""
    for pages in pdf.pages:
        page_text=pages.extract_text()
        if page_text:
            text=text+page_text+"\n"
    return text


def read_docx(file_path:Path):
    document=Document(file_path)
    text=""
    for paragraphs in document.paragraphs:
        if paragraphs.text.strip():
            text=text+paragraphs.text +"\n"
    for tables in document.tables:
        for row in tables.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text=text+cell.text.strip()+"\n"
    return text
def read_resume(file_path:Path):
    file_path=Path(file_path)
    if file_path.suffix.lower()==".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower()==".docx":
        return read_docx(file_path)
    else:
        return None
